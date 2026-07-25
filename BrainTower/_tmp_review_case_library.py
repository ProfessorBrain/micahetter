from __future__ import annotations

import json
import sys
from pathlib import Path

from docx import Document


def paragraph_payload(paragraph):
    text = paragraph.text.strip()
    if not text:
        return None
    return {
        "style": paragraph.style.name if paragraph.style else "",
        "text": text,
    }


def table_payload(table):
    return [
        [cell.text.strip() for cell in row.cells]
        for row in table.rows
    ]


def main() -> None:
    source = Path(sys.argv[1])
    output_path = Path(sys.argv[2]) if len(sys.argv) > 2 else None
    document = Document(source)

    paragraphs = [
        payload
        for paragraph in document.paragraphs
        if (payload := paragraph_payload(paragraph)) is not None
    ]
    tables = [table_payload(table) for table in document.tables]
    headings = [
        paragraph
        for paragraph in paragraphs
        if paragraph["style"].lower().startswith("heading")
        or paragraph["style"].lower() in {"title", "subtitle"}
    ]

    output = {
        "source": str(source),
        "paragraph_count": len(paragraphs),
        "table_count": len(tables),
        "section_count": len(document.sections),
        "headings": headings,
        "paragraphs": paragraphs,
        "tables": tables,
    }
    if output_path:
        output_path.write_text(
            json.dumps(output, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
    else:
        json.dump(output, sys.stdout, ensure_ascii=False, indent=2)


if __name__ == "__main__":
    main()
